from docx import Document
import openpyxl

wb = openpyxl.Workbook()
ws = wb.worksheets[0]

document = Document('123.docx')
allpar = document.paragraphs
alltab = document.tables
s = len(allpar)
l = len(alltab)
# print(s)
print(len(alltab))

for i in range(0,l,2):                        
    ps=alltab[i].rows[5].cells[1].text
    hmy=alltab[i].rows[1].cells[1].text
    my=alltab[i].rows[2].cells[1].text
    jbmy=alltab[i].rows[3].cells[1].text
    bmy=alltab[i].rows[4].cells[1].text
    s = allpar[i].text.split('(')
    # print(s[0][-3:])l
    # print(s[1][0:2])
    
    ws.cell(i/2+1, 1, eval(ps))
    ws.cell(i/2+1, 2, eval(hmy))
    ws.cell(i/2+1, 3, eval(my))
    ws.cell(i/2+1, 4, eval(jbmy))
    ws.cell(i/2+1, 5, eval(bmy))
    ws.cell(i/2+1, 6, s[0])
    ws.cell(i/2+1, 7, s[1])
    
wb.save('123.xlsx')
wb.close()